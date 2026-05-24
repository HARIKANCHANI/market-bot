#!/usr/bin/env python3
"""
Convert DOCUMENTATION_SUMMARY.md to Word document (.docx)
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
except ImportError:
    print("❌ Required libraries not found!")
    print("\nPlease install:")
    print("  pip install python-docx")
    print("\nThen run this script again.")
    exit(1)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    print(f"📄 Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print(f"📝 Creating Word document...")
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    
    for line in lines:
        # Skip empty lines at start
        if not line.strip() and len(doc.paragraphs) == 0:
            continue
            
        # H1 Headers (# )
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # H2 Headers (## )
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            
        # H3 Headers (### )
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            
        # Horizontal rule (---)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            
        # Code blocks (```)
        elif line.strip().startswith('```'):
            continue  # Skip code block markers
            
        # Bullet points (- or *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown bold (**text**)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Remove markdown code (`text`)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            
        # Tables (|)
        elif '|' in line and line.strip().startswith('|'):
            # Simple table handling - just add as paragraph
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
        # Regular text
        elif line.strip():
            text = line
            # Remove markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Remove markdown code
            text = re.sub(r'`(.*?)`', r'\1', text)
            # Remove markdown links [text](url)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            p = doc.add_paragraph(text)
        else:
            doc.add_paragraph()  # Empty line
    
    print(f"💾 Saving to {docx_file}...")
    doc.save(docx_file)
    print(f"✅ Successfully created {docx_file}")

if __name__ == "__main__":
    import sys

    # Default to DATA_FLOW_DETAILED.md
    input_file = "DATA_FLOW_DETAILED.md"
    output_file = "DATA_FLOW_DETAILED.docx"

    # Allow command line arguments
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]

    print("=" * 60)
    print("📄 Markdown to Word Converter")
    print("=" * 60)
    print()

    try:
        convert_markdown_to_docx(input_file, output_file)
        print()
        print("🎉 Conversion complete!")
        print(f"📁 Output file: {output_file}")
        print()
        print("You can now open the Word document:")
        print(f"  start {output_file}")
        print()
    except FileNotFoundError:
        print(f"❌ Error: {input_file} not found!")
        print("Make sure the file exists in the current directory.")
    except Exception as e:
        print(f"❌ Error: {e}")
